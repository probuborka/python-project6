"""Модуль для генерации отчета о структуре файлов и папок на жестком диске."""
import argparse
import json
import zipfile
from abc import ABC, abstractmethod
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document

def check_path(path):
    """Проверка пути к папке."""
    path = Path(path)
    
    if not path.exists():
        raise ValueError(f"Путь '{path}' не существует")

def check_file(report):
    """Проверка файла."""
    check_path(report)

    path = Path(report)
    
    if path.is_dir():
        raise ValueError(f"'{path}' не является файлом")

def tree_to_strings(data, prefix='', is_root=True, string=None):
    """Рекурсивно переводим дерево в список строк."""
    if is_root:
        is_root = False
    
    # сортируем (папки сначала, затем файлы)
    items = []
    for key, value in data.items():
        items.append((key, value))
    
    # cортируем: сначала папки, потом файлы
    items.sort(key=lambda x: (0 if x[1].get('type') == 'folder' or x[1].get('type') == 'zip' else 1, x[0]))
    
    # Обрабатываем каждый элемент
    for _, (key, value) in enumerate(items): 
        prefix_tmp = '      '       
        current_prefix = prefix + prefix_tmp
        
        item_type = value.get('type', 'file')
        
        # Форматируем размер
        size = value.get('size', 0)
        size_str = f'{size} bytes'
        
        # Форматируем дату
        modif_date = value.get('modif_date', '')
        try:
            date_obj = datetime.strptime(modif_date, '%Y-%m-%d %H:%M:%S')
            date_str = date_obj.strftime('%d.%m.%Y %H:%M')
        except (ValueError, TypeError):
            date_str = modif_date
        
        # Выводим текущий элемент
        string.append(f"{prefix}{prefix_tmp} {value.get('name', key)} ({value.get('type', '')}/{size_str}/{date_str})")
        
        # Рекурсивно обрабатываем вложенные элементы
        if ( item_type == 'folder' or item_type == 'zip' ):
            # Ищем вложенные элементы
            nested_data = {}
            for k, v in value.items():
                if k not in ['name', 'type', 'path', 'size', 'modif_date']:
                    nested_data[k] = v
            
            if nested_data:
                tree_to_strings(nested_data, current_prefix, False, string)

class DirectoryTree:
    """Получение структуры файлов и папок на жестком диске."""
    def scan(self, path):
        """Сканирует указанную директорию и возвращает ее структуру."""
        def zip_tree(zip_file, structure):
            """Обработка ZIP."""
            with zip_file as zipf:
                # Получаем файлы и папки
                for file_info in zipf.infolist():
                    path = Path(file_info.filename)
                    extension = path.suffix.lower()  
                    path_parts = path.parts
                    
                    # Построение дерева
                    current_level = structure
                    
                    # переход на узел (папку)
                    for i, part in enumerate(path_parts):
                        if i < len(path_parts)-1:
                            key = f"{'dir'}_{part}"
                            current_level = current_level[key]
                    
                    # создание узла (папки) или файла
                    if file_info.is_dir():
                        key = f"{'dir'}_{part}"
                        current_level[key] = { 
                                            'name': part,
                                            'type': 'folder',
                                            'path': f'{str(path)}',
                                            'size': file_info.file_size,
                                            'modif_date': datetime(*file_info.date_time).strftime('%Y-%m-%d %H:%M:%S')
                                            }
                    else:
                        # если этой zip обрабатываем его
                        if extension == '.zip':
                            key = f"{'zip'}_{part}"
                            zfiledata = BytesIO(zipf.read(part))
                            current_level[key] = zip_tree(
                                zipfile.ZipFile(zfiledata),
                                { 
                                    'name': part,
                                    'type': 'zip',
                                    'path': f'{str(path)}',
                                    'size': file_info.file_size,
                                    'modif_date': datetime(*file_info.date_time).strftime('%Y-%m-%d %H:%M:%S')
                                }
                            )
                        else:
                        # файл
                            key = f"{'fol'}_{part}"
                            current_level[key] = { 
                                                'name': part, 
                                                'type': 'file',
                                                'path': f'{str(path)}',
                                                'size': file_info.file_size,
                                                'modif_date': datetime(*file_info.date_time).strftime('%Y-%m-%d %H:%M:%S')
                                                }

            return structure

        def folder_tree(path):
            """Обработка DIR."""
            structure = {}

            path = Path(path)
            
            # Получаем все файлы и папки
            for file_path in path.rglob('*'):
                extension = file_path.suffix.lower() 
                relative_path = file_path.relative_to(path)
                path_parts = relative_path.parts

                # Построение дерева
                current_level = structure

                # переход на узел (папку)
                for i, part in enumerate(path_parts):
                    if i < len(path_parts)-1:
                        key = f"{'dir'}_{part}"
                        current_level = current_level[key]

                # создание узла (папки) или файла
                if file_path.is_dir():
                    key = f"{'dir'}_{part}"
                    current_level[key] = { 
                                            'name': part,   
                                            'type': 'folder',
                                            'path': f'{str(relative_path)}',
                                            'size': file_path.stat().st_size,
                                            'modif_date': datetime.fromtimestamp(file_path.stat().st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                                        }
                else:
                    # если этой zip обрабатываем его
                    if extension == '.zip':
                        key = f"{'zip'}_{part}"
                        
                        current_level[key] = zip_tree( 
                            zipfile.ZipFile(file_path, 'r'),
                            { 
                                'name': part,
                                'type': 'zip',
                                'path': f'{str(relative_path)}',
                                'size': file_path.stat().st_size,
                                'modif_date': datetime.fromtimestamp(file_path.stat().st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                            }
                        )
                    else:
                    # файл
                        key = f"{'fol'}_{part}"
                        current_level[key] = { 
                                                'name': part,
                                                'type': 'file',
                                                'path': f'{str(relative_path)}',
                                                'size': file_path.stat().st_size,
                                                'modif_date': datetime.fromtimestamp(file_path.stat().st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                                            }
            
            return structure
        
        check_path(path)

        return folder_tree(path) 

class Saver(ABC):
    """Абстрактный класс для всех форматов сохранения."""
    @abstractmethod
    def save(self, data, filename):
        """Сохранение."""
        pass

class JSONSaver(Saver):
    """Сохранение в json."""
    def save(self, data, file):
        """Сохранение."""
        check_file(file)

        json_output = json.dumps(data, indent=2, ensure_ascii=False)
        file_path = Path(file)
        with file_path.open('w', encoding='utf-8') as f:
            f.write(json_output)
        return 'Данные сохранены в .json'

class DOCXSaver(Saver):
    """Сохранение в doc."""
    def save(self, data, file):
        """Сохранение."""
        check_file(file)

        string = []
        tree_to_strings(data=data, is_root=True, string=string)

        doc = Document()
        doc.add_heading('Пример .Docx', 0)
        for v in string:
            doc.add_paragraph(v)
        doc.save(file)
        return 'Данные сохранены в .docx'

class PDFSaver(Saver):
    """Сохранение в pdf."""
    def save(self, data, file):
        """Сохранение."""
        pass

class SaverFactory:
    """Фабрика для создания объектов сохранения."""
    @staticmethod
    def create_saver(format_type):
        """Создает объект Saver для указанного формата."""
        savers = {
            '.json': JSONSaver,
            '.docx' : DOCXSaver,
            '.pdf' : PDFSaver
        }
        
        saver = savers.get(format_type.lower())

        if not saver:
            raise ValueError(f'Неподдерживаемый формат: {format_type}')
        
        return saver()

def main():
    """Точка входа."""
    parser = argparse.ArgumentParser(description='Анализ структуры файлов и папок')
    parser.add_argument('--path', type=str, required=True, 
                       help='Путь к анализируемой папке ')
    
    parser.add_argument('--report', type=str, required=True, 
                       help='Путь к отчету')

    args = parser.parse_args()

    # args.path = '/home/user/dev/python/project/python-project6/folder_first'
    # args.report = '/home/user/dev/python/project/python-project6/report/example.docx'

    format_type = Path(args.report).suffix.lower()
    
    try: 
        dir_tree = DirectoryTree()
        tree = dir_tree.scan(args.path)
    except ValueError as e:
        print(f'Ошибка: {e}')
        return

    try:  
        factory = SaverFactory()
        saver   = factory.create_saver(format_type)
        result  = saver.save(tree, args.report)

        print(f'{result}')
    except ValueError as e:
        print(f'Ошибка: {e}')
        return
    
    
if __name__ == '__main__':

    main()